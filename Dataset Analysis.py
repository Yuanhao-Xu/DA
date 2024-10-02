import xgboost as xgb
import os
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from sklearn.preprocessing import StandardScaler
from docx import Document
from docx.shared import Inches


# Function to check for missing values and save the result
def check_missing_values(data, folder_name):
    missing_values = data.isnull().sum()
    missing_values_file = os.path.join(folder_name, 'missing_values_report.csv')
    missing_values.to_csv(missing_values_file)
    return missing_values, missing_values_file


# Function to generate descriptive statistics
def descriptive_statistics(data, folder_name):
    descriptive_stats = data.describe()
    descriptive_stats_file = os.path.join(folder_name, 'descriptive_statistics_report.csv')
    descriptive_stats.to_csv(descriptive_stats_file)
    return descriptive_stats, descriptive_stats_file


# Function to calculate correlation matrix and save it
def correlation_analysis(data, folder_name):
    correlation_matrix = data.corr()
    correlation_matrix_file = os.path.join(folder_name, 'correlation_matrix.csv')
    correlation_matrix.to_csv(correlation_matrix_file)
    return correlation_matrix, correlation_matrix_file


# Function to plot histograms
def plot_histograms(data, folder_name):
    plt.figure(figsize=(15, 10))
    data.hist(bins=30, figsize=(15, 10))
    plt.tight_layout()
    histogram_file = os.path.join(folder_name, 'data_histograms.png')
    plt.savefig(histogram_file)
    plt.close()  # Close the plot to save memory
    return histogram_file


# Function to plot correlation heatmap
def plot_correlation_heatmap(correlation_matrix, folder_name):
    plt.figure(figsize=(10, 8))
    sns.heatmap(correlation_matrix, annot=True, cmap='coolwarm', fmt='.2f')
    plt.title('Correlation Heatmap')
    plt.tight_layout()
    heatmap_file = os.path.join(folder_name, 'correlation_heatmap.png')
    plt.savefig(heatmap_file)
    plt.close()  # Close the plot to save memory
    return heatmap_file


# Modified function to calculate and plot feature importances using XGBoost
def feature_importance_analysis_xgb(data, folder_name):
    X = data.drop(data.columns[-1], axis=1)  # Assuming the last column is the target
    y = data[data.columns[-1]]

    # Initialize XGBoost Regressor
    model = xgb.XGBRegressor(random_state=42)

    # Fit the model
    model.fit(X, y)

    # Extract feature importances
    feature_importances = model.feature_importances_
    feature_importance_df = pd.DataFrame({
        'Feature': X.columns,
        'Importance': feature_importances
    }).sort_values(by='Importance', ascending=False)

    # Plot feature importances
    plt.figure(figsize=(10, 6))
    sns.barplot(x='Importance', y='Feature', data=feature_importance_df)
    plt.title('XGBoost Feature Importances')
    plt.tight_layout()

    # Save the plot
    feature_importance_file = os.path.join(folder_name, 'xgb_feature_importance.png')
    plt.savefig(feature_importance_file)
    plt.close()  # Close the plot to save memory

    # Save feature importance report as CSV
    feature_importance_report = os.path.join(folder_name, 'xgb_feature_importances_report.csv')
    feature_importance_df.to_csv(feature_importance_report, index=False)

    return feature_importance_df, feature_importance_file, feature_importance_report


# Function to save results into a Word document
def save_to_word(missing_values, descriptive_stats, correlation_matrix, folder_name, histogram_file, heatmap_file,
                 feature_importance_file):
    doc = Document()
    doc.add_heading('Dataset Analysis Report', 0)

    doc.add_heading('Missing Values Analysis', level=1)
    doc.add_paragraph(missing_values.to_string())

    doc.add_heading('Descriptive Statistics', level=1)
    doc.add_paragraph(descriptive_stats.to_string())

    doc.add_heading('Correlation Matrix', level=1)
    doc.add_paragraph(correlation_matrix.to_string())

    doc.add_heading('Histograms of Features', level=1)
    doc.add_picture(histogram_file, width=Inches(5))

    doc.add_heading('Correlation Heatmap', level=1)
    doc.add_picture(heatmap_file, width=Inches(5))

    doc.add_heading('Feature Importances', level=1)
    doc.add_picture(feature_importance_file, width=Inches(5))

    doc_file = os.path.join(folder_name, 'dataset_analysis_report.docx')
    doc.save(doc_file)
    return doc_file


# Main function to perform the entire analysis for each dataset
def perform_analysis_for_datasets(paths):
    # Create the main folder for storing all reports
    main_folder = '/mnt/data/Dataset_Analysis_Reports'
    os.makedirs(main_folder, exist_ok=True)

    # Loop over each dataset in the paths dictionary
    for dataset_name, dataset_path in paths.items():
        # Create a subfolder for each dataset
        dataset_folder = os.path.join(main_folder, dataset_name)
        os.makedirs(dataset_folder, exist_ok=True)

        # Load the dataset
        data = pd.read_csv(dataset_path)

        # Perform all analysis steps for this dataset
        missing_values, missing_values_file = check_missing_values(data, dataset_folder)
        descriptive_stats, descriptive_stats_file = descriptive_statistics(data, dataset_folder)
        correlation_matrix, correlation_matrix_file = correlation_analysis(data, dataset_folder)
        histogram_file = plot_histograms(data, dataset_folder)
        heatmap_file = plot_correlation_heatmap(correlation_matrix, dataset_folder)
        feature_importance_df, feature_importance_file, feature_importance_report = feature_importance_analysis_xgb(
            data,
            dataset_folder)

        # Save results to Word document
        doc_file = save_to_word(missing_values, descriptive_stats, correlation_matrix, dataset_folder, histogram_file,
                                heatmap_file, feature_importance_file)

        print(f"Analysis report for {dataset_name} saved in {dataset_folder}")


# Define the paths dictionary (this is the one you provided)
paths = {
    "UCI_concrete": "Dataset/concrete/concrete_data.csv",
    "BFRC_cs": "Dataset/BFRC/data_cs.csv",
    "BFRC_fs": "Dataset/BFRC/data_fs.csv",
    "BFRC_sts": "Dataset/BFRC/data_sts.csv",
    "Pullout_fmax": "Dataset/pullout/dataset_fmax.csv",
    "Pullout_ifss": "Dataset/pullout/dataset_ifss.csv",
    "CST": "Dataset/Concrete_Slump_Test/data.csv",
    "UHPC_cs": "Dataset/uhpc/Compressive_strength.csv",
    "UHPC_fs": "Dataset/uhpc/Flexural_strength.csv",
    "UHPC_mss": "Dataset/uhpc/Mini_slump_spread.csv",
    "UHPC_porosity": "Dataset/uhpc/Porosity.csv",
    "ENB2012_HL": "Dataset/ENB2012/data1.csv",
    "ENB2012_CL": "Dataset/ENB2012/data2.csv",
}

# Run the analysis for all datasets
perform_analysis_for_datasets(paths)
