import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Pt, Inches
import os

class R2Analyzer:
    def __init__(self, data_dicts, r2_thresholds, total_samples, addendum_init, addendum_size):
        """
        Initialize the class.
        :param data_dicts: Dictionary of datasets.
        :param r2_thresholds: List of R2 thresholds.
        :param total_samples: Total number of samples per dataset.
        :param addendum_init: Initial number of samples.
        :param addendum_size: Number of samples added per cycle.
        """
        self.data_dicts = data_dicts
        self.r2_thresholds = r2_thresholds
        self.total_samples = total_samples
        self.addendum_init = addendum_init
        self.addendum_size = addendum_size

    def find_r2_proportions(self):
        """
        Returns a dictionary with strategy names. Each key contains another dictionary
        where the key is the dataset and the value is two lists:
        1. Sample proportion for reaching each R2 threshold.
        """
        strategy_dict = {}
        for dataset_name, strategy_data in self.data_dicts.items():
            for strategy, r2_values in strategy_data.items():
                if strategy not in strategy_dict:
                    strategy_dict[strategy] = {}
                sample_counts = []
                proportions = []
                for threshold in self.r2_thresholds:
                    for i, r2_value in enumerate(r2_values):
                        if r2_value >= threshold:
                            samples_used = self.addendum_init + i * self.addendum_size
                            proportion_used = samples_used / self.total_samples
                            sample_counts.append(samples_used)
                            proportions.append(proportion_used)
                            break
                    else:
                        sample_counts.append(None)
                        proportions.append(None)
                strategy_dict[strategy][dataset_name] = proportions
        return strategy_dict

    def plot_r2_proportions(self, proportions_dict, image_folder):
        """
        Plot sample proportions for each strategy and dataset, saving the images.
        All plots have the same x-axis range [0, 1].
        :param proportions_dict: Dictionary with sample proportions for each dataset.
        :param image_folder: Folder to save the images.
        """
        if not os.path.exists(image_folder):
            os.makedirs(image_folder)

        for strategy, dataset_proportions in proportions_dict.items():
            plt.figure(figsize=(8, 6))
            for dataset_name, proportions in dataset_proportions.items():
                plt.plot(proportions, self.r2_thresholds, marker='o', label=dataset_name)

            plt.xlabel('Sample Proportion')
            plt.ylabel('R2 Thresholds')
            plt.title(f'{strategy} - Proportion of Samples for Each Threshold')
            plt.legend()
            plt.grid(True)

            # Set y-axis ticks to show only thresholds
            plt.yticks(self.r2_thresholds)

            # Fix x-axis range to [0, 1]
            plt.xlim([0, 1])

            # Save the image
            image_path = os.path.join(image_folder, f'{strategy}_r2_proportions.png')
            plt.savefig(image_path)
            plt.close()

    def save_images_and_tables_to_word(self, result, image_folder, word_file_name='dataset_analysis_report.docx'):
        """
        Save all images and proportion tables into a Word document with formatting.
        :param result: Dictionary with strategy names, dataset proportions.
        :param image_folder: Folder containing the images.
        :param word_file_name: Name of the generated Word document.
        """
        word_path = os.path.join(image_folder, word_file_name)
        doc = Document()

        # Set document title
        doc.add_heading('R2 Proportions Analysis', 0)

        # Insert images
        for strategy in result.keys():
            image_path = os.path.join(image_folder, f'{strategy}_r2_proportions.png')
            if os.path.exists(image_path):
                doc.add_heading(f'{strategy} Strategy - R2 Proportions', level=1)
                doc.add_picture(image_path, width=Inches(5.5))

        # Insert tables
        for strategy, dataset_proportions in result.items():
            doc.add_heading(f'{strategy} Strategy - Proportion Table', level=1)
            table = doc.add_table(rows=1, cols=len(self.r2_thresholds) + 1)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Dataset'
            for i, threshold in enumerate(self.r2_thresholds):
                hdr_cells[i + 1].text = f'R2 ≥ {threshold}'

            for dataset_name, proportions in dataset_proportions.items():
                row_cells = table.add_row().cells
                row_cells[0].text = dataset_name
                for i, proportion in enumerate(proportions):
                    row_cells[i + 1].text = f'{proportion:.2f}' if proportion is not None else 'N/A'

        # Save document
        doc.save(word_path)
