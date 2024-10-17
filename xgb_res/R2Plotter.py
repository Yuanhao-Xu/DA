from docx import Document
from docx.shared import Pt
import os
import matplotlib.pyplot as plt

class R2Plotter:
    def __init__(self, dataset_name, data_dict, addendum_init, addendum_size):
        """
        Initialize class with dataset name, data dictionary, initial samples, and sample increment.
        """
        self.dataset_name = dataset_name
        self.data_dict = data_dict
        self.addendum_init = addendum_init
        self.addendum_size = addendum_size

        # Create a folder named after the dataset
        if not os.path.exists(self.dataset_name):
            os.makedirs(self.dataset_name)

    def plot_all(self):
        """
        Generate plots for all strategies and save as PNG.
        """
        plt.figure(figsize=(16, 8))  # Set figure width

        for strategy, r2_scores in self.data_dict.items():
            cycles = range(1, len(r2_scores) + 1)
            # Calculate data volume per cycle
            data_volume = [self.addendum_init + (cycle - 1) * self.addendum_size for cycle in cycles]
            # Plot each strategy
            plt.plot(data_volume, r2_scores, label=strategy, linestyle='-', linewidth=2, alpha=0.8)

        # Set title and labels
        plt.title(f'R^2 Scores of Different Strategies on {self.dataset_name} Dataset', fontsize=24)
        plt.xlabel('Data Volume', fontsize=22)
        plt.ylabel('R^2 Score', fontsize=22)

        # Set legend location
        plt.legend(loc='lower right', fontsize=18)
        plt.grid(True)

        # Add gridlines
        plt.grid(which='both', linestyle='--', linewidth=0.5)

        # Adjust layout
        plt.tight_layout()

        # Save as PNG
        file_name = f"All_Strategies_on_{self.dataset_name}.png"
        plt.savefig(os.path.join(self.dataset_name, file_name))
        plt.close()  # Close to save memory

    def plot_vs_RS(self):
        """
        Generate comparison plot between each strategy and RS, and save as PNG.
        """
        for strategy, r2_scores in self.data_dict.items():
            if strategy == 'RS':
                continue  # Skip RS itself

            plt.figure(figsize=(16, 8))  # Set figure width

            cycles = range(1, len(r2_scores) + 1)
            # Calculate data volume per cycle
            data_volume = [self.addendum_init + (cycle - 1) * self.addendum_size for cycle in cycles]

            # Plot RS and other strategies
            plt.plot(data_volume, self.data_dict['RS'], label='RS', linestyle='-', linewidth=2, alpha=0.8)
            plt.plot(data_volume, r2_scores, label=strategy, linestyle='-', linewidth=2, alpha=0.8)

            # Set title and labels
            plt.title(f'R^2 Scores Comparison between {strategy} and RS on {self.dataset_name} Dataset', fontsize=24)
            plt.xlabel('Data Volume', fontsize=22)
            plt.ylabel('R^2 Score', fontsize=22)

            # Show legend
            plt.legend(loc='lower right', fontsize=18)
            plt.grid(True)

            # Add gridlines
            plt.grid(which='both', linestyle='--', linewidth=0.5)

            # Adjust layout
            plt.tight_layout()

            # Save as PNG
            file_name = f"{strategy}_vs_RS_on_{self.dataset_name}.png"
            plt.savefig(os.path.join(self.dataset_name, file_name))
            plt.close()  # Close to save memory

    def get_r2_samples(self, r2_thresholds, total_samples):
        """
        Return number and proportion of samples when each strategy reaches R2 thresholds.
        """
        result = {}

        # Loop through each strategy
        for strategy, r2_scores in self.data_dict.items():
            strategy_result = {}
            samples = [self.addendum_init + (cycle - 1) * self.addendum_size for cycle in range(1, len(r2_scores) + 1)]

            # Find first occurrence for each R2 threshold
            for r2_threshold in r2_thresholds:
                threshold_met = False
                for idx, r2_score in enumerate(r2_scores):
                    if r2_score >= r2_threshold:
                        sample_count = samples[idx]
                        proportion = sample_count / total_samples
                        strategy_result[r2_threshold] = {
                            'sample_count': sample_count,
                            'proportion': proportion
                        }
                        threshold_met = True
                        break

                # If not reached, set to N/A
                if not threshold_met:
                    strategy_result[r2_threshold] = {
                        'sample_count': 'N/A',
                        'proportion': 'N/A'
                    }

            result[strategy] = strategy_result

            # Print results
            print(f"Strategy: {strategy}")
            for r2_value, info in strategy_result.items():
                if info['sample_count'] == 'N/A':
                    print(f"  R2 >= {r2_value}: N/A")
                else:
                    print(f"  R2 >= {r2_value}: Sample Count = {info['sample_count']}, Proportion = {info['proportion']:.2%}")

        return result

    def save_report(self, result, r2_thresholds, image_folder, filename=None):
        """
        Save R2 report to Word, including tables and images.
        """
        # Set default filename if not provided
        if filename is None:
            filename = os.path.join(self.dataset_name, 'r2_report.docx')

        # Create Word document
        doc = Document()
        doc.add_heading(f'R2 Samples Report for {len(r2_thresholds)} Thresholds', 0)

        # Create table with header
        table = doc.add_table(rows=1, cols=len(r2_thresholds) + 1)

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Strategy'

        # Generate columns for each R2 threshold
        for i, r2_value in enumerate(r2_thresholds):
            hdr_cells[i + 1].text = f'R2 >= {r2_value}'

        # Fill in data for each strategy
        for strategy, data in result.items():
            row_cells = table.add_row().cells
            row_cells[0].text = strategy
            for i, r2_value in enumerate(r2_thresholds):
                if data[r2_value]['sample_count'] == 'N/A':
                    row_cells[i + 1].text = 'N/A'
                else:
                    sample_count = data[r2_value]['sample_count']
                    proportion = f"{data[r2_value]['proportion']:.2%}"
                    row_cells[i + 1].text = f"{sample_count} ({proportion})"

        # Set font size for table
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.strip()
                cell.paragraphs[0].runs[0].font.size = Pt(10)

        # Insert images
        for image_file in os.listdir(image_folder):
            if image_file.endswith('.png'):
                doc.add_page_break()
                doc.add_paragraph(image_file)
                doc.add_picture(os.path.join(image_folder, image_file), width=Pt(400))

        # Save the document
        doc.save(filename)
        print(f"Report saved as {filename}")
